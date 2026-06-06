from __future__ import annotations

import json
import re
import sqlite3
import textwrap
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from urllib.error import URLError
from urllib.request import urlopen

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
PDF_SOURCE = Path(r"C:\Users\p venkata naveen\Downloads\quant_forex_V10 Research.pdf")
OUT_DOCX = DOCS / "quant_forex_V10_Complete_Operator_Manual.docx"
OUT_SNAPSHOT = DOCS / "quant_forex_V10_Complete_Operator_Manual_sources.json"


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def api_get(path: str):
    try:
        with urlopen(f"http://127.0.0.1:8000{path}", timeout=5) as r:
            return json.loads(r.read().decode("utf-8"))
    except Exception:
        return None


def extract_pdf_text(path: Path) -> dict:
    info = {
        "path": str(path),
        "exists": path.exists(),
        "page_count": 0,
        "sample_text": "",
        "pages": [],
        "error": None,
    }
    if not path.exists():
        return info
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        chunks = []
        for page in reader.pages:
            chunks.append(page.extract_text() or "")
        text = "\n".join(chunks)
        info["page_count"] = len(reader.pages)
        info["sample_text"] = text[:12000]
        info["pages"] = [
            {"page": i + 1, "text": re.sub(r"\s+", " ", (page.extract_text() or "")).strip()[:1200]}
            for i, page in enumerate(reader.pages)
        ]
    except Exception as exc:
        info["error"] = repr(exc)
    return info


def db_schema_summary(database_py: Path) -> list[dict[str, str]]:
    text = database_py.read_text(encoding="utf-8")
    tables = []
    for match in re.finditer(r"CREATE TABLE IF NOT EXISTS\s+([a-zA-Z0-9_]+)\s*\((.*?)\n\s*\)", text, re.S):
        name = match.group(1)
        body = re.sub(r"\s+", " ", match.group(2)).strip()
        columns = []
        for raw in body.split(","):
            col = raw.strip()
            if not col or col.upper().startswith(("PRIMARY", "FOREIGN", "UNIQUE", "CHECK")):
                continue
            columns.append(col.split()[0])
        tables.append({"table": name, "columns": ", ".join(columns[:18]) + ("..." if len(columns) > 18 else "")})
    return tables


def endpoint_summary(app_py: Path) -> list[dict[str, str]]:
    text = app_py.read_text(encoding="utf-8")
    rows = []
    for method, path, fn in re.findall(r"@app\.(get|post|put|delete)\(\"([^\"]+)\".*?\ndef\s+([a-zA-Z0-9_]+)", text, re.S):
        rows.append({"method": method.upper(), "path": path, "function": fn})
    return rows


def clean(s, limit: int | None = None) -> str:
    if s is None:
        return ""
    if isinstance(s, (dict, list)):
        s = json.dumps(s, ensure_ascii=False)
    s = str(s).replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"\s+", " ", s).strip()
    if limit and len(s) > limit:
        return s[: limit - 3].rstrip() + "..."
    return s


def add_page_number_note(doc: Document):
    section = doc.sections[-1]
    footer = section.footer.paragraphs[0]
    footer.text = "quant_forex_V10 Complete Operator Manual - Research/backtest only, no live order execution"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(90, 103, 125)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(31, 78, 121)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], clean(val, 900))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str], style: str = "List Bullet"):
    for item in items:
        if not item:
            continue
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(1)
        p.add_run(clean(item, 1000))


def add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(textwrap.dedent(text).strip())
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(35, 45, 65)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(31, 78, 121)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(31, 78, 121)
        else:
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(36, 54, 79)
    return p


def para(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    return p


def build_sources():
    regimes = load_json(ROOT / "backend" / "common" / "config" / "regimes.json")
    strategies = load_json(ROOT / "backend" / "common" / "config" / "strategies.json")
    modifiers = load_json(ROOT / "backend" / "common" / "config" / "modifiers.json")
    formulas = load_json(ROOT / "backend" / "common" / "config" / "formulas.json")
    thresholds = load_json(ROOT / "backend" / "common" / "config" / "thresholds.json")
    market = load_json(ROOT / "backend" / "common" / "config" / "market.json")

    api = {
        "health": api_get("/api/health"),
        "mode_presets": api_get("/api/reference/mode-presets"),
        "data_sources": api_get("/api/data-sources"),
        "api_structure": api_get("/api/reference/api-structure"),
    }
    sources = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "repo": str(ROOT),
        "source_pdf": extract_pdf_text(PDF_SOURCE),
        "regime_count": len(regimes),
        "strategy_count": len(strategies),
        "modifier_count": len(modifiers),
        "formula_count": len(formulas),
        "regimes": regimes,
        "strategies": strategies,
        "modifiers": modifiers,
        "formulas": formulas,
        "thresholds": thresholds,
        "market": market,
        "tables": db_schema_summary(ROOT / "backend" / "database.py"),
        "endpoints": endpoint_summary(ROOT / "backend" / "app.py"),
        "api": api,
    }
    OUT_SNAPSHOT.write_text(json.dumps(sources, indent=2, ensure_ascii=False), encoding="utf-8")
    return sources


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"
    add_page_number_note(doc)
    return doc


def add_cover(doc: Document, sources: dict):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("quant_forex_V10\nComplete Operator Manual")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Regime research, strategy testing, pattern engine, MT5 validation, UI usage, saved data, and practical forex workflow")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(68, 84, 106)

    doc.add_paragraph()
    add_table(
        doc,
        ["Item", "Current Document Source"],
        [
            ["App", "Local FastAPI + one-page UI at http://127.0.0.1:8000/"],
            ["Health scope", clean((sources.get("api", {}).get("health") or {}).get("scope"))],
            ["Regimes", f"{sources['regime_count']} regimes from backend/common/config/regimes.json"],
            ["Strategies", f"{sources['strategy_count']} strategies from backend/common/config/strategies.json"],
            ["PDF reviewed", f"{sources['source_pdf']['path']} ({sources['source_pdf']['page_count']} pages extracted)" if sources["source_pdf"]["exists"] else "PDF not found"],
            ["Generated", sources["generated_at"]],
        ],
        [1.8, 5.7],
    )
    para(
        doc,
        "Important: This system is a research, reference, backtest, validation, and MT5 report analysis console. It does not place live broker orders. Use it to decide what deserves demo or live manual review, not as a guarantee of profit.",
    )
    doc.add_page_break()


def add_toc(doc: Document):
    heading(doc, "How To Read This Manual", 1)
    para(doc, "This document is written as a working desk manual. Read the first three sections before using the app. Use the later sections as reference while backtesting.")
    add_table(
        doc,
        ["Section", "Use It For"],
        [
            ["1. Fast Operating Map", "Know the complete app workflow in one screen."],
            ["2. UI Controls", "Understand every input, checkbox, mode, button, and output panel."],
            ["3. Forex Logic", "Understand regimes, strategies, patterns, modifiers, costs, and validation."],
            ["4. Workflows", "Run discovery, strict validation, final approval, JSON request testing, optimizer, and MT5 report import."],
            ["5. Saved Data", "Know what is stored in SQLite and how it should appear in UI."],
            ["6. Reference Tables", "Look up all regimes, strategies, formulas, endpoints, and database tables."],
            ["7. Covered vs Missing", "Know what is implemented and what still needs external setup or more research."],
        ],
        [1.8, 5.7],
    )


def classify_pdf_page(text: str) -> tuple[str, str]:
    lower = text.lower()
    if "research controls" in lower:
        return (
            "Research Controls and top setup",
            "Use these pages to choose symbol, timeframe, regime, strategy, mode preset, macro evidence, data source, strict filters, pattern engine, and statistical regime controls before running anything.",
        )
    if "regime research lab" in lower:
        return (
            "Regime Research Lab",
            "Use this when you want one selected regime to expose its editable thresholds, mapped strategies, optimizer ranges, and candidate validation settings.",
        )
    if "monthly regime sweep" in lower or "saved monthly sweeps" in lower:
        return (
            "Monthly regime sweep",
            "Use this to test whether a regime/strategy works across months instead of only across the full period average.",
        )
    if "experiment builder" in lower or "a/b" in lower:
        return (
            "A/B experiments and calibration",
            "Use variants to compare one controlled change at a time, then apply calibration profiles only after the variant improves evidence without overfitting.",
        )
    if "parity" in lower or "real tick validation" in lower or "mt5 tester" in lower:
        return (
            "MT5 tester, parity, and real-tick lane",
            "Use this for final execution realism: prepare MT5 config, run tester models externally, import reports, and compare Python vs MT5 by deterministic keys.",
        )
    if "validation cockpit" in lower or "out-of-sample" in lower or "walk-forward" in lower:
        return (
            "Validation cockpit and OOS/WF",
            "Use this after a candidate looks good. It must survive unseen data, walk-forward windows, Monte Carlo stress, and optional portfolio breadth.",
        )
    if "optimizer grid" in lower or "monte carlo" in lower:
        return (
            "Optimizer and Monte Carlo",
            "Use optimizer to create candidates, then Monte Carlo to study drawdown tails, losing streaks, and probability of loss. Optimizer alone does not approve a system.",
        )
    if "saved backtest" in lower or "current regime" in lower:
        return (
            "Saved runs and current regime",
            "Use saved runs/favorites to reload prior evidence. Current Regime shows the latest detected state and why it is blocked or active.",
        )
    if "data health" in lower or "feature summary" in lower:
        return (
            "Data health and features",
            "Use this to confirm bars, missing data, duplicates, warmup, ADX/ER/VWAP/statistical features, and whether the data is good enough to trust.",
        )
    if "regime cards" in lower:
        return (
            "Regime cards",
            "Use cards as the reference library and drilldown entry. Click a regime to filter trades, performance, modifier impact, skipped reasons, and approval checklist.",
        )
    if "pattern performance" in lower or "trade list" in lower or "skipped" in lower:
        return (
            "Performance and trade diagnostics",
            "Use these panels to see what worked, what failed, which patterns helped, which setups were skipped, and whether stops are too tight via MAE/MFE.",
        )
    if "llm" in lower or "ollama" in lower:
        return (
            "LLM reviewer and libraries",
            "Use local Ollama or deterministic fallback to review latest evidence; then inspect strategy, modifier, and formula libraries for exact definitions.",
        )
    if "formula reference" in lower:
        return (
            "Formula reference and API examples",
            "Use this to verify the actual formulas and request shapes behind the UI. This is the debugging reference when a result does not make sense.",
        )
    return (
        "App snapshot page",
        "Use this page as visual context for the one-page research console. Match the visible block to the detailed section in this manual.",
    )


def add_pdf_page_guide(doc: Document, sources: dict):
    heading(doc, "Source PDF Page-By-Page Usage Guide", 1)
    pdf = sources.get("source_pdf") or {}
    para(
        doc,
        f"The supplied PDF snapshot has {pdf.get('page_count', 0)} pages. The table below maps the visible PDF pages to what the app section is for and how to operate it.",
    )
    rows = []
    for page in pdf.get("pages", []):
        title, use = classify_pdf_page(page.get("text", ""))
        rows.append([str(page.get("page", "")), title, use, clean(page.get("text", ""), 260)])
    if rows:
        for i in range(0, len(rows), 12):
            add_table(doc, ["PDF Page", "Visible Section", "How To Use It", "Visible Text Clue"], rows[i : i + 12], [0.75, 1.7, 2.9, 2.15])
    else:
        para(doc, "The PDF text could not be extracted, so the manual uses the live app and backend configs as the primary source.")


def add_fast_map(doc: Document):
    heading(doc, "1. Fast Operating Map", 1)
    para(doc, "The app should be used as a research pipeline, not as a single signal button. The correct mental model is:")
    add_code(
        doc,
        """
        Raw MT5/SQLite candle data
        -> feature engine
        -> data-quality check
        -> regime detection
        -> strategy eligibility
        -> pattern engine and modifiers
        -> alpha score and final score
        -> backtest trades
        -> OOS / walk-forward / Monte Carlo / optimizer
        -> MT5 real-tick report import and model comparison
        -> LLM/deterministic review
        -> final approval / watchlist / reject
        -> saved run and favorites in UI
        """,
    )
    add_table(
        doc,
        ["Layer", "What Works Here", "Why It Matters"],
        [
            ["Regime", "Classifies market state such as trend, range, news, stress, session, macro, cost, VWAP, and institutional patterns.", "A strategy that works in R01 can fail badly in R03/R10/R50."],
            ["Strategy", "Defines entry, SL/TP style, risk/reward, and direction for one regime.", "Prevents using every setup in every market."],
            ["Pattern", "Optional confluence such as FVG, OB, BOS, MSS, liquidity pool, round number, VWAP/MVWAP/session VWAP.", "Adds measurable confirmation. It should improve results or be removed."],
            ["Modifier", "Session, spread, trend weakening, MTF conflict, news, rollover, sentiment, macro, cost, and quality adjustments.", "Separates valid setups from poor execution conditions."],
            ["Validation", "Backtest, OOS, walk-forward, Monte Carlo, model comparison, MT5 report import, parity checks.", "Tells whether an edge is robust or just fitted to one sample."],
            ["Storage", "Runs, trades, validation results, macro data, MT5 reports, favorites.", "Lets you compare historical research instead of repeating guesswork."],
        ],
        [1.2, 3.1, 3.1],
    )
    para(doc, "Short answer to the trading question: with this setup you can research semi-manual ideas, but you should only move to demo review after local backtest, OOS, walk-forward, Monte Carlo, and MT5 real-tick validation agree. Funded/live trading needs additional broker execution discipline, risk kill-switches, and manual trade governance.")


def add_ui_controls(doc: Document):
    heading(doc, "2. UI Controls And How To Use Them", 1)
    heading(doc, "2.1 Top Setup Controls", 2)
    add_table(
        doc,
        ["Control", "What It Does", "How To Use It"],
        [
            ["Symbol", "Market to test, for example EURUSD, GBPUSD, USDJPY, XAUUSD.", "Start with one liquid major pair. Do not mix symbols until one pair is stable."],
            ["Timeframe", "Bar timeframe for the research logic.", "M15 is a practical first setting. M5 needs stricter spread/tick validation. H1 needs larger sample periods."],
            ["Start / End Date", "Backtest window.", "Use at least 6-12 months for initial discovery, then separate train/test windows."],
            ["Regime", "ALL or one selected Rxx regime.", "Use ALL for discovery, then single regime for diagnosis."],
            ["Strategy", "ALL or one selected strategy.", "When a regime is selected, UI should show only that regime's allowed strategies."],
            ["Risk %", "Research risk per trade.", "Use 0.25-0.75% for funded-style research. Use 1% only for comparison."],
            ["RR", "Default reward/risk target if strategy does not override it.", "Use 1.5 for mean reversion; 2.0 for trend/breakout; test 2.5/3.0 only after sample size is strong."],
            ["Initial Equity", "Starting account used for P/L numbers.", "Use your intended prop account size for realistic drawdown context."],
            ["Sentiment / USD / Risk / Central Bank Bias", "Manual macro bias inputs or evidence-mode context.", "Manual values should not activate macro regimes by themselves unless evidence confidence is present."],
        ],
        [1.35, 2.9, 3.1],
    )
    heading(doc, "2.2 Mode Presets", 2)
    add_table(
        doc,
        ["Preset", "Use When", "Expected Settings"],
        [
            ["Discovery", "You are searching for candidates.", "Looser filters, score_only pattern mode, broader sessions, min alpha around 5-6, 1-Min OHLC or candle research."],
            ["Strict Validation", "A candidate looked promising and needs cleaner proof.", "Harder spread/session filters, min alpha 7-8, stricter regime validation, OOS/WF/MC required."],
            ["Final Approval", "A candidate may be considered for demo/manual trading review.", "Every Tick Based On Real Ticks, hard filters, min alpha 8+, spread max 65, real MT5 report/model comparison, final approval gate."],
        ],
        [1.5, 2.7, 3.3],
    )
    heading(doc, "2.3 Pattern Engine Controls", 2)
    add_table(
        doc,
        ["Control", "Meaning", "Practical Rule"],
        [
            ["Pattern engine", "Master switch for pattern confirmation.", "Turn ON for institutional-style confluence testing. Turn OFF to see raw regime/strategy edge."],
            ["ICT", "Umbrella switch for measurable ICT-style concepts.", "Use only as measurable features; do not treat as magic."],
            ["FVG", "Three-candle imbalance zone.", "Useful after displacement. Track size in ATR, age, fill %, and invalidation."],
            ["Order blocks", "Last opposite candle before displacement/BOS.", "Test freshness and mitigation. Avoid old, repeatedly tapped zones."],
            ["BOS", "Break of swing structure.", "Require close beyond swing by an ATR buffer, not only a wick."],
            ["MSS", "Market structure shift against the prior trend.", "Stronger when preceded by liquidity sweep and displacement."],
            ["Liquidity pools", "Equal highs/lows, Asia high/low, PDH/PDL, PWH/PWL.", "Use as level context; entry still needs reclaim/rejection."],
            ["Round numbers", "00/50 pip zones and symbol-specific steps.", "Helpful for FX behavior, but must be tested per pair."],
            ["VWAP", "Session fair value using tick volume.", "Use mean reversion in R36 and acceptance/continuation in R46."],
            ["MVWAP", "Rolling VWAP such as 20/50 periods.", "Use as dynamic fair value in trends."],
            ["Session VWAP", "VWAP reset by Asia/London/New York.", "Use for session-specific behavior, not all-day assumptions."],
            ["Min pattern score", "Minimum pattern confluence needed.", "2 is discovery. 3+ is stricter. Hard minimum can reduce trades sharply."],
            ["Pattern score mode", "score_only or minimum_required/hard mode.", "score_only lets trades pass with penalty/bonus. Hard mode blocks weak pattern setups."],
        ],
        [1.55, 2.85, 3.1],
    )
    heading(doc, "2.4 MT5 Backtest Model", 2)
    add_table(
        doc,
        ["Model", "What It Means", "When To Use"],
        [
            ["Candle Close", "Very rough bar-close approximation.", "Fast sanity check only. Do not trust tight SL/TP results."],
            ["1-Min OHLC", "Uses one-minute OHLC approximation.", "Best for first research pass across many combinations."],
            ["Every Tick", "MT5 generated tick sequence.", "Validation when strategy is spread/slippage sensitive."],
            ["Every Tick Based On Real Ticks", "Uses broker's real tick history where available.", "Final approval for scalping, tight stops, VWAP, SL/TP order, and funded-style validation."],
        ],
        [1.7, 2.8, 3.0],
    )
    heading(doc, "2.5 Strict Filters", 2)
    add_table(
        doc,
        ["Control", "Why It Exists", "Recommended Starting Value"],
        [
            ["Strict regime validation", "Prevents fuzzy confidence from activating a clean regime while key rules fail.", "ON for strict validation and final approval."],
            ["Reject trend weakening", "Blocks R01/R02 trades when ADX/ER slopes deteriorate.", "ON for clean trend systems."],
            ["Reject low ER", "Protects trend strategies from noisy chop.", "ON for clean trend, OFF only when testing range logic."],
            ["Reject ADX outside band", "Stops R01/R02 from mixing with high-vol/exhaustion regimes.", "ON for regime purity."],
            ["Reject MTF conflict score > 0", "Blocks HTF/LTF disagreement for clean trend continuation.", "ON for R01/R02/R44, optional for trap regimes."],
            ["Min alpha score", "Minimum total quality score before trade is allowed.", "5 discovery, 7 strict, 8-9 final approval."],
            ["Max spread percentile", "Blocks costly execution states.", "70 discovery, 65 final, 90 only for stress tagging."],
            ["Killzone mode", "score_only or hard_filter session behavior.", "hard_filter for final validation."],
            ["Spread mode", "score_only or hard_filter spread behavior.", "hard_filter for final validation."],
        ],
        [1.65, 3.25, 2.6],
    )


def add_forex_logic(doc: Document):
    heading(doc, "3. Forex Logic: What The App Is Measuring", 1)
    para(doc, "A regime is the market environment. A strategy is the trade model. A pattern is optional confirmation. A modifier is context that can improve or damage the trade. A filter is a hard yes/no gate. A cost model adjusts whether the trade is realistic after spread, slippage, commission, and tick sequence.")
    add_table(
        doc,
        ["Concept", "Forex Meaning", "Correct Use"],
        [
            ["ADX", "Directional strength, not direction.", "Trend regimes need moderate/strong ADX; range regimes need low ADX."],
            ["ER", "Kaufman Efficiency Ratio: clean movement vs noisy path.", "Trend systems need ER high enough; low ER means chop."],
            ["ATR percentile", "Current volatility relative to recent history.", "Low vol supports compression/drift; high vol supports breakout/stress/exhaustion."],
            ["Spread percentile", "Current spread relative to recent spread history.", "High spread can erase edge, especially M5/M15 and tight stops."],
            ["Session", "Asia/London/NewYork/Overlap/Rollover behavior.", "London/NY/Overlap usually best for liquidity; rollover is defensive/no-trade."],
            ["HTF/LTF bias", "Higher timeframe vs lower timeframe direction.", "Alignment helps trend. Conflict can be a trap regime but hurts clean trend."],
            ["VWAP", "Tick-volume weighted intraday fair value.", "FX has no centralized volume, so MT5 tick volume is a proxy."],
            ["FVG/OB/BOS/MSS", "OHLC approximations of imbalance, supply/demand, and structure.", "Must be defined numerically and tested, not traded as narrative."],
            ["COT/macro/cross-pair", "External evidence for institutional positioning and broad currency regimes.", "Useful for bias, but lower frequency and not a standalone entry."],
        ],
        [1.4, 3.0, 3.1],
    )
    heading(doc, "3.1 Pattern Formulas In Practical Terms", 2)
    add_table(
        doc,
        ["Pattern", "Measurable Definition", "What To Record Per Trade"],
        [
            ["Bullish FVG", "Candle 1 high < Candle 3 low; size = zone height / ATR.", "zone_low, zone_high, size_atr, age_bars, fill_percent, invalidated."],
            ["Bearish FVG", "Candle 1 low > Candle 3 high.", "zone_low, zone_high, size_atr, age_bars, fill_percent, invalidated."],
            ["Bullish OB", "Last bearish candle before bullish displacement/BOS.", "ob_low, ob_high, age_bars, mitigated, displacement_atr."],
            ["Bearish OB", "Last bullish candle before bearish displacement/BOS.", "ob_low, ob_high, age_bars, mitigated, displacement_atr."],
            ["BOS", "Close beyond swing high/low by ATR buffer.", "level, close_distance_atr, direction, swing_age."],
            ["MSS", "Sweep against previous trend then break structure the other way.", "sweep_level, structure_level, displacement, confirmation bar."],
            ["Liquidity pool", "Equal highs/lows, session high/low, PDH/PDL, PWH/PWL.", "pool_type, level, distance_atr, swept_or_rejected."],
            ["Round number", "00/50 pip zone with symbol-specific step.", "round_level, distance_atr, rejection/acceptance."],
            ["VWAP/MVWAP", "sum(typical_price * tick_volume) / sum(tick_volume).", "distance_from_vwap_atr, reclaim, acceptance, reversion target."],
        ],
        [1.35, 3.2, 3.0],
    )
    heading(doc, "3.2 Why Some Backtests Fail", 2)
    add_bullets(
        doc,
        [
            "Regime contamination: a clean trend regime allowed low ER, high spread, or ADX outside the allowed band.",
            "Execution costs: spread/slippage consumed the expected R, especially with stops below 0.75-1.0 ATR.",
            "Overfit pattern combinations: one exact FVG/VWAP/OB setup worked in-sample but failed OOS.",
            "Session mismatch: profitable London behavior was mixed with Asia or OffSession trades.",
            "Insufficient sample: 20-30 trades can look good but is not enough for funded-style confidence.",
            "MT5 model mismatch: candle/1-min result did not survive real ticks.",
            "Macro inconsistency: USD/risk/central-bank regime was selected manually without evidence confirmation.",
        ],
    )


def add_workflows(doc: Document):
    heading(doc, "4. Practical Workflows", 1)
    workflows = [
        (
            "4.1 First Discovery Run",
            [
                "Set Symbol = EURUSD, Timeframe = M15, Date range = at least 6 months.",
                "Set Regime = ALL and Strategy = ALL.",
                "Set Mode = Discovery.",
                "Pattern engine ON, Pattern score mode = score_only, Min pattern score = 2.",
                "MT5 model = 1-Min OHLC for speed.",
                "Run Backtest and read regime, strategy, pattern, session, month, and combination tables.",
                "Star only candidates with enough trades, positive expectancy, PF > 1.10, and understandable reasons.",
            ],
        ),
        (
            "4.2 Single Regime Diagnosis",
            [
                "Select one regime, for example R36 VWAP / Mean-Reversion Normal-Vol.",
                "The strategy dropdown should show VW1, VW2, VW3 only.",
                "Run each strategy separately.",
                "Compare pattern performance: VWAP alone, FVG+VWAP, BOS+VWAP, liquidity sweep+VWAP.",
                "Look at skipped setup reasons. If many are blocked by spread/session, do not loosen immediately; first confirm market quality.",
            ],
        ),
        (
            "4.3 Strict Validation",
            [
                "Use the same candidate from discovery.",
                "Set Mode = Strict Validation.",
                "Turn ON strict regime validation, reject trend weakening, reject low ER, reject ADX outside band, reject MTF conflict.",
                "Set Min alpha score = 7 or 8, Max spread percentile = 70 or 65.",
                "Run OOS, walk-forward, and Monte Carlo.",
                "If trades fall to zero, lower one threshold at a time. Do not loosen all filters together.",
            ],
        ),
        (
            "4.4 Final Approval Before Demo Review",
            [
                "Run the exact same settings using Every Tick Based On Real Ticks.",
                "Import the MT5 report and model comparison.",
                "Run final approval review.",
                "Candidate should pass local backtest, OOS, walk-forward, Monte Carlo, MT5 real-tick comparison, and anti-overfit checks.",
                "If any one layer fails, keep it as research/watchlist only.",
            ],
        ),
        (
            "4.5 JSON Input / Output Console",
            [
                "Open the Regime + Strategy JSON Input / Output section.",
                "Select a regime/strategy card.",
                "Edit payload values such as risk, RR, filters, pattern options, MT5 test model, and strict controls.",
                "Click Send Request.",
                "Right side should show the full response: sent payload, regime details, strategy details, config, results, skipped reasons, pattern details, and validation status.",
                "Use this like Postman for one setup at a time.",
            ],
        ),
        (
            "4.6 Optimizer / Permutation Research",
            [
                "Choose one regime and one strategy first, not the whole universe.",
                "Define ranges: min_alpha_score, max_spread_percentile, RR, pattern score, FVG size, FVG age, session filters, VWAP distance.",
                "Run optimizer grid.",
                "Save only candidates that pass sample size and basic PF/expectancy gates.",
                "Immediately run OOS/WF/MC. Optimizer rank alone is not approval.",
            ],
        ),
    ]
    for title, steps in workflows:
        heading(doc, title, 2)
        add_bullets(doc, [f"{i + 1}. {step}" for i, step in enumerate(steps)], "List Number")


def add_saved_data(doc: Document, sources: dict):
    heading(doc, "5. What Is Saved In The Database And How UI Should Show It", 1)
    para(doc, "The app should make research history visible. A user should not need to remember which run was good. Saved runs and favorites are part of the research workflow.")
    add_table(
        doc,
        ["Database Table", "What It Stores", "UI Should Show"],
        [
            ["candles", "Symbol, timeframe, OHLC, tick volume, spread and timestamp.", "Data health, last candle time, count by symbol/timeframe."],
            ["features", "Calculated indicators per candle.", "Feature completeness, latest ADX/ER/ATR/VWAP/spread/regime inputs."],
            ["feature_cache", "Cached feature calculation payloads.", "Cache status, stale/fresh status, recalculation button."],
            ["backtest_runs", "Backtest summary, request payload, performance tables, skipped reasons.", "Saved runs list, star/favorite first, reload run details."],
            ["backtest_trades", "Trade-level entry/exit, regime, strategy, R, profit, context, pattern/alpha/final score.", "Trade table with filters and detailed expanded rows."],
            ["mt5_report_imports", "Imported MT5 report summary and metadata.", "MT5 reports list, model used, source file, pass/fail status."],
            ["mt5_report_deals", "MT5 deal/trade rows.", "Deal-level comparison to Python backtest and execution diagnostics."],
            ["macro_data", "COT, cross-pair, macro, rates, risk, external evidence rows.", "Macro evidence panel with confidence and date freshness."],
            ["ab_experiments", "Baseline vs variants experiment results.", "A/B table with deltas and decision gates."],
            ["validation_runs", "OOS, WF, MC, optimizer, portfolio, model comparison, final approval summaries.", "Validation history sorted by starred first, then newest."],
            ["monthly_regime_sweep_runs", "Monthly sweep research summary.", "Month-by-month stability panel."],
            ["monthly_regime_sweep_candidates", "Candidate combinations from monthly sweeps.", "Best candidates with validation status."],
            ["favorites", "Starred backtests, validation runs, experiments, features, MT5 reports.", "Star icon on every major result section; starred items shown first."],
        ],
        [1.8, 2.75, 2.95],
    )
    heading(doc, "5.1 Full Schema Snapshot", 2)
    add_table(doc, ["Table", "Detected Columns"], [[t["table"], t["columns"]] for t in sources["tables"]], [1.8, 5.7])


def add_coverage(doc: Document):
    heading(doc, "6. Covered, Partly Covered, Not Covered", 1)
    add_table(
        doc,
        ["Area", "Status", "What This Means"],
        [
            ["50-regime reference", "Covered", "Regimes R01-R50 exist in config and UI/API reference should show them."],
            ["Strategy map", "Covered", "Strategies are mapped to regimes; UI should filter strategy options by selected regime."],
            ["Pattern engine", "Covered/needs continuous calibration", "FVG, OB, BOS, MSS, liquidity, round numbers, VWAP/MVWAP/session VWAP are research features."],
            ["Strict controls", "Covered", "Strict validation, alpha, spread/session, trend weakening, low ER, ADX band, MTF conflict can be UI controlled."],
            ["OOS/WF/MC", "Covered", "Validation engines exist; results need to be saved and read in UI history."],
            ["Optimizer", "Covered", "Can rank parameter combinations; should not approve alone."],
            ["MT5 Strategy Tester automation", "Partly covered", "Config/runner/import/parity modules exist, but real MT5 terminal setup and reports are external dependencies."],
            ["MT5 live order execution", "Not covered", "No live or demo order placement should be assumed."],
            ["True interbank order book", "Not covered", "Retail MT5 does not provide centralized spot FX order book."],
            ["Bloomberg/news/options positioning", "Not covered", "Requires paid institutional feeds or separate integrations."],
            ["Funded account readiness", "Not automatic", "Only after robust validation, demo review, and risk governance."],
        ],
        [1.8, 1.5, 4.2],
    )


def add_reference_tables(doc: Document, sources: dict):
    heading(doc, "7. Regime Reference: R01-R50", 1)
    para(doc, "Use this as the master market-state map. Each regime tells you when a strategy is allowed, blocked, reduced, or watchlist only.")
    regimes = sources["regimes"]
    for chunk_start in range(0, len(regimes), 10):
        rows = []
        for r in regimes[chunk_start : chunk_start + 10]:
            rows.append(
                [
                    r["regime_id"],
                    r["regime_name"],
                    r.get("direction", ""),
                    ", ".join(r.get("allowed_strategies", [])),
                    clean(r.get("meaning", ""), 220),
                ]
            )
        add_table(doc, ["ID", "Regime", "Direction", "Allowed Strategies", "Meaning"], rows, [0.6, 1.8, 0.8, 1.6, 2.7])

    heading(doc, "8. Detailed Regime Cards", 1)
    for r in regimes:
        heading(doc, f"{r['regime_id']} - {r['regime_name']}", 2)
        add_table(
            doc,
            ["Field", "Value"],
            [
                ["Meaning", r.get("meaning", "")],
                ["Direction", r.get("direction", "")],
                ["Allowed strategies", ", ".join(r.get("allowed_strategies", []))],
                ["Blocked strategies", ", ".join(r.get("blocked_strategies", []))],
                ["Risk", clean(r.get("risk", {}), 450)],
                ["Conditions", " | ".join(r.get("conditions", []))],
                ["Rules", " | ".join(r.get("rules", []))],
            ],
            [1.4, 6.1],
        )

    heading(doc, "9. Strategy Reference", 1)
    by_regime = defaultdict(list)
    for s in sources["strategies"]:
        by_regime[str(s.get("regime", ""))].append(s)
    for regime, strategies in sorted(by_regime.items()):
        heading(doc, f"Strategies for {regime}", 2)
        add_table(
            doc,
            ["Strategy", "Name", "Direction", "Category", "Default RR"],
            [[s["strategy_id"], s["strategy_name"], s.get("direction", ""), s.get("category", ""), str(s.get("default_rr", ""))] for s in strategies],
            [0.9, 3.1, 1.0, 1.6, 0.9],
        )


def add_formulas_and_endpoints(doc: Document, sources: dict):
    heading(doc, "10. Formula And Feature Reference", 1)
    formulas = sources.get("formulas", [])
    if isinstance(formulas, dict):
        rows = [[k, clean(v, 500)] for k, v in formulas.items()]
    else:
        rows = []
        for item in formulas:
            if isinstance(item, dict):
                rows.append([item.get("name") or item.get("formula_id") or item.get("id") or "formula", clean(item, 500)])
            else:
                rows.append(["formula", clean(item, 500)])
    if rows:
        for i in range(0, len(rows), 20):
            add_table(doc, ["Formula", "Meaning / Definition"], rows[i : i + 20], [1.7, 5.8])
    else:
        para(doc, "No formula config rows were detected, but backend/common/base/formulas.py contains executable feature calculations.")

    heading(doc, "11. API Endpoint Reference", 1)
    endpoint_rows = [[e["method"], e["path"], e["function"]] for e in sources["endpoints"]]
    for i in range(0, len(endpoint_rows), 28):
        add_table(doc, ["Method", "Path", "Backend Function"], endpoint_rows[i : i + 28], [0.8, 4.1, 2.6])


def add_examples(doc: Document):
    heading(doc, "12. Example Requests And Interpretations", 1)
    heading(doc, "12.1 R01 Clean Bullish Trend With T1", 2)
    add_code(
        doc,
        """
        {
          "symbol": "EURUSD",
          "timeframe": "M15",
          "start_date": "2025-11-17",
          "end_date": "2026-05-17",
          "regime_filter": "R01",
          "strategy_filter": "T1",
          "risk_percent": 0.5,
          "rr": 2,
          "patterns": {
            "use_fvg": true,
            "use_bos": true,
            "use_vwap": true,
            "min_pattern_score": 2,
            "pattern_score_mode": "score_only"
          },
          "filters": {
            "strict_regime_validation": true,
            "reject_trend_weakening": true,
            "reject_low_er": true,
            "reject_adx_outside_band": true,
            "reject_mtf_conflict": true,
            "min_alpha_score": 8,
            "max_spread_percentile": 65,
            "killzone_mode": "hard_filter",
            "spread_filter_mode": "hard_filter"
          }
        }
        """,
    )
    para(doc, "Interpretation: this is a strict clean-trend test. If trades become too few, reduce only one filter at a time: first max spread 65->70, then min alpha 8->7, then pattern hard minimum to score_only. Do not disable regime purity first.")
    heading(doc, "12.2 R36 VWAP Mean Reversion", 2)
    add_bullets(
        doc,
        [
            "Use when ADX <= 18, ER <= 0.25, normal volatility, no news, and price is at least 1.5 ATR away from session VWAP.",
            "VW1 is short from high VWAP stretch; VW2 is long from low VWAP stretch; VW3 tests failed extension.",
            "Good combinations to test: VWAP + liquidity pool, VWAP + round number, VWAP + failed BOS, VWAP + session boundary.",
            "Bad combinations: VWAP fade during R44 Trend Day or R46 VWAP Trend Acceptance.",
        ],
    )
    heading(doc, "12.3 If OOS Is Weak", 2)
    add_table(
        doc,
        ["Symptom", "Likely Cause", "Next Action"],
        [
            ["Great in-sample, weak OOS", "Overfit thresholds or one market phase.", "Widen train/test windows, lower parameter count, test different months and symbols."],
            ["Good PF but few trades", "Insufficient sample.", "Keep as watchlist; do not approve. Expand date range or adjacent pairs."],
            ["Good candle result, bad real ticks", "Tight stops/cost sensitivity.", "Increase ATR stop buffer, use hard spread filter, retest real ticks."],
            ["Many skipped setups", "Filters too strict or data quality gaps.", "Read skipped reasons. Fix data first; only then calibrate thresholds."],
            ["Losses cluster in OffSession", "Session edge mismatch.", "Use hard killzone filter or separate OffSession regime."],
        ],
        [1.9, 2.75, 2.85],
    )


def add_proceed_plan(doc: Document):
    heading(doc, "13. How You Should Proceed From Here", 1)
    add_table(
        doc,
        ["Phase", "Goal", "Action", "Promotion Gate"],
        [
            ["Research", "Find possible edges.", "ALL regimes/strategies, Discovery mode, 1-Min OHLC, score_only patterns.", "Enough trades, PF > 1.10, expectancy > 0, reasons make sense."],
            ["Diagnosis", "Understand where it works/fails.", "Single regime + strategy, pattern combinations, session/month/modifier breakdown.", "Weak zones identified and blocked without curve fitting."],
            ["Strict Validation", "Check robustness.", "Strict controls, OOS, WF, MC, anti-overfit.", "Passes OOS/WF/MC with acceptable drawdown."],
            ["MT5 Validation", "Check execution realism.", "Every Tick Based On Real Ticks, import report, model comparison.", "Real tick result is close enough to research result."],
            ["Demo Review", "Human process check.", "Manual/demo only, journal every signal and missed trade.", "At least several weeks without process errors."],
            ["Funded/Live Consideration", "Risk-managed deployment.", "Only after written rules, daily loss cap, no revenge trading, broker spread checks.", "Never use if validation or psychology is unstable."],
        ],
        [1.35, 1.8, 2.6, 1.75],
    )
    para(doc, "The safest posture: use the app to reject weak ideas quickly and promote only a small number of robust, explainable setups. More trades are not the target. Cleaner evidence is the target.")


def build_document():
    sources = build_sources()
    doc = setup_doc()
    add_cover(doc, sources)
    add_toc(doc)
    add_pdf_page_guide(doc, sources)
    add_fast_map(doc)
    add_ui_controls(doc)
    add_forex_logic(doc)
    add_workflows(doc)
    add_saved_data(doc, sources)
    add_coverage(doc)
    add_reference_tables(doc, sources)
    add_formulas_and_endpoints(doc, sources)
    add_examples(doc)
    add_proceed_plan(doc)
    doc.save(OUT_DOCX)
    print(json.dumps({"docx": str(OUT_DOCX), "snapshot": str(OUT_SNAPSHOT), "regimes": sources["regime_count"], "strategies": sources["strategy_count"]}, indent=2))


if __name__ == "__main__":
    build_document()
